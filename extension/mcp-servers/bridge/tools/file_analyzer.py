"""Universal File Analyzer - analyze any uploaded file by type"""
import os, json, base64, time
from pathlib import Path
from datetime import datetime
from bridge.utils import _markdown_header, _markdown_footer
from bridge.config import DZ_SESSION_FILE

def _get_file_info(path: str) -> dict:
    """Get file metadata"""
    p = Path(path)
    stat = p.stat()
    ext = p.suffix.lower()
    size_kb = stat.st_size / 1024
    
    # Determine type category
    image_exts = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg', '.ico'}
    doc_exts = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt'}
    text_exts = {'.txt', '.md', '.rst', '.json', '.yaml', '.yml', '.csv', '.xml', '.toml', '.ini', '.cfg'}
    code_exts = {'.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.cpp', '.c', '.h', '.hpp', 
                 '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.sh', '.bat', '.ps1'}
    archive_exts = {'.zip', '.tar', '.gz', '.rar', '.7z'}
    
    if ext in image_exts:
        ftype = "image"
    elif ext in doc_exts:
        ftype = "document"
    elif ext in text_exts:
        ftype = "text"
    elif ext in code_exts:
        ftype = "code"
    elif ext in archive_exts:
        ftype = "archive"
    else:
        ftype = "binary"
    
    return {
        "path": str(p.absolute()),
        "name": p.name,
        "ext": ext,
        "type": ftype,
        "size_kb": round(size_kb, 1),
        "size_mb": round(size_kb / 1024, 2),
        "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }

def _read_text_file(path: str, max_chars: int = 50000) -> str:
    """Read text file with encoding detection"""
    encodings = ["utf-8", "euc-kr", "cp949", "latin-1", "shift-jis"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                text = f.read(max_chars)
            if len(text) >= max_chars:
                text += f"\n\n...[truncated at {max_chars} chars]"
            return text, enc
        except UnicodeDecodeError:
            continue
    return "", "unknown"


def _encode_image_as_safe_data_uri(image_path: str, max_dim: int = 1024, quality: int = 85) -> str:
    """이미지를 안전한 Data URI로 인코딩.
    
    Pillow로 이미지를 로드하여 max_dim 이하로 리사이징한 후,
    JPEG으로 압축(품질 85)하고 전체 Base64 인코딩을 수행한다.
    문자열 슬라이싱을 사용하지 않으므로 이미지 손상이 발생하지 않는다.
    """
    try:
        from PIL import Image
        import io
        
        img = Image.open(image_path)
        
        # RGBA → RGB 변환 (JPEG은 알파채널 미지원)
        if img.mode in ('RGBA', 'P', 'LA'):
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = rgb_img
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 리사이징 (max_dim 이하로, 비율 유지)
        w, h = img.size
        largest = max(w, h)
        if largest > max_dim:
            scale = max_dim / largest
            new_w = int(w * scale)
            new_h = int(h * scale)
            img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        
        # JPEG 압축 → Base64
        buffer = io.BytesIO()
        img.save(buffer, format='JPEG', quality=quality, optimize=True)
        img_b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        return f"data:image/jpeg;base64,{img_b64}"
    except ImportError:
        # Pillow 미설치 시: 원본 이미지를 4자 정렬로 안전하게 자르기 (fallback)
        try:
            with open(image_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode('utf-8')
            safe_len = (200000 // 4) * 4
            ext = os.path.splitext(image_path)[1].lower().lstrip('.') or 'png'
            if ext == 'jpg': ext = 'jpeg'
            return f"data:image/{ext};base64,{img_b64[:safe_len]}"
        except Exception:
            return ""
    except Exception:
        return ""


def analyze_file(file_path: str) -> str:
    """파일 분석 메인 함수
    
    Args:
        file_path: 분석할 파일 경로
    Returns:
        마크다운 형식 분석 보고서
    """
    path = os.path.expanduser(file_path)
    if not os.path.exists(path):
        return (_markdown_header("File Error", "❌")
                + f"**File not found:** `{path}`\n"
                + _markdown_footer())
    
    info = _get_file_info(path)
    
    lines = []
    lines.append(f"## 📄 File Analysis: {info['name']}")
    lines.append("")
    lines.append(f"- **Type:** {info['type']} ({info['ext']})")
    lines.append(f"- **Size:** {info['size_kb']} KB ({info['size_mb']} MB)")
    lines.append(f"- **Path:** `{info['path']}`")
    lines.append(f"- **Modified:** {info['modified']}")
    lines.append("")
    
    # For images: 통합 Image Pipeline 실행
    if info['type'] == 'image':
        # 안전한 Data URI 생성 (Pillow 리사이징)
        safe_uri = _encode_image_as_safe_data_uri(path, max_dim=1024, quality=85)
        if safe_uri:
            lines.append(f"![uploaded image]({safe_uri})")
            lines.append("")
        lines.append("---")
        lines.append("## 🔬 Image Analysis Pipeline")
        lines.append("")

        # ── Phase 1: SSA (Statistical Spatial Aggregator) ──
        try:
            from bridge.tools.ssa import _analyze_image, _imread_korean_safe, _summarize_ssa_results
            import cv2, numpy as np
            img_raw = _imread_korean_safe(path)
            if img_raw is not None:
                orig_h, orig_w = img_raw.shape[:2]
                target_w = 640
                target_h = int(orig_h * (target_w / orig_w))
                img_resized = cv2.resize(img_raw, (target_w, target_h))
                ssa_report = _analyze_image(img_resized, detail="full", orig_w=orig_w, orig_h=orig_h)
                ssa_summary = _summarize_ssa_results(ssa_report)
                if ssa_summary:
                    lines.append(ssa_summary)
                lines.append(ssa_report)
                lines.append("")
            else:
                lines.append("### 📊 SSA: Cannot read image for spatial analysis\n")
        except ImportError:
            lines.append("### 📊 SSA: OpenCV not available (install opencv-python)\n")
        except Exception as e:
            lines.append(f"### 📊 SSA: Analysis failed ({e})\n")

        # ── Phase 2: OCR 텍스트 추출 ──
        try:
            from bridge.ocr_engine import OcrEngine
            ocr_engine = OcrEngine()
            if ocr_engine.is_available():
                ocr_result = ocr_engine.ocr(path, lang="auto", detail="quick")
                ocr_md = OcrEngine.ocr_to_markdown(ocr_result)
                lines.append(ocr_md)
                lines.append("")
            else:
                lines.append("### 📝 OCR: Not available (install Tesseract or PaddleOCR)\n")
        except ImportError:
            lines.append("### 📝 OCR: Module not loaded\n")
        except Exception as e:
            lines.append(f"### 📝 OCR: Failed ({e})\n")

        # ── Phase 3: MiniCPM-V Vision ──
        try:
            from bridge.vision.minicpm import describe_image, is_available
            if is_available():
                desc = describe_image(path)
                if desc:
                    lines.append("### 🤖 Vision Analysis (MiniCPM-V)")
                    lines.append(desc)
                    lines.append("")
        except Exception:
            pass
    
    # For text/code files: read content
    if info['type'] in ('text', 'code'):
        content, encoding = _read_text_file(path)
        lines.append(f"### 📝 Content ({encoding})")
        lines.append(f"```{info['ext'].lstrip('.')}")
        lines.append(content[:10000])
        if len(content) > 10000:
            lines.append("...[truncated]")
        lines.append("```")
        lines.append("")
    
    # For documents: try text extraction
    if info['type'] == 'document':
        if info['ext'] == '.pdf':
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(path)
                lines.append(f"### 📑 PDF Content ({doc.page_count} pages)")
                has_text = False
                for i, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        has_text = True
                        lines.append(f"\n**Page {i+1}:**")
                        lines.append(f"```\n{text[:2000]}\n```")
                    if i >= 5:
                        lines.append("*...more pages available*")
                        break
                doc.close()
                if not has_text:
                    # 스캔 문서: 이미지 변환 → OCR/MiniCPM 파이프라인
                    _analyze_pdf_as_image(path, lines)
            except ImportError:
                lines.append("⚠️ PyMuPDF not installed. Run: `pip install PyMuPDF`")
            except Exception as e:
                lines.append(f"⚠️ PDF read error: {e}")
        
        elif info['ext'] == '.docx':
            try:
                import docx
                d = docx.Document(path)
                text = "\n".join(p.text for p in d.paragraphs)
                lines.append(f"### 📝 DOCX Content ({len(d.paragraphs)} paragraphs)")
                lines.append(f"```\n{text[:5000]}\n```")
            except ImportError:
                lines.append("⚠️ python-docx not installed. Run: `pip install python-docx`")
            except Exception as e:
                lines.append(f"⚠️ DOCX read error: {e}")
    
    # For binary files: show hex preview
    if info['type'] == 'binary':
        try:
            with open(path, "rb") as f:
                data = f.read(512)
            hex_str = data.hex()
            # Format as hex dump
            dump_lines = []
            for i in range(0, len(hex_str), 32):
                addr = i // 2
                hex_part = " ".join(hex_str[j:j+2] for j in range(i, min(i+32, len(hex_str)), 2))
                ascii_part = "".join(chr(int(hex_str[j:j+2], 16)) if 32 <= int(hex_str[j:j+2], 16) < 127 else "." for j in range(i, min(i+32, len(hex_str)), 2))
                dump_lines.append(f"{addr:08x}  {hex_part:<48}  {ascii_part}")
            lines.append("### 🔤 Hex Preview (first 512 bytes)")
            lines.append("```")
            lines.extend(dump_lines[:16])
            lines.append("```")
        except Exception:
            pass
    
    # Add summary
    lines.insert(1, f"*{info['size_kb']}KB {info['type']} file analyzed*")
    
    return "\n".join(lines)


def _analyze_pdf_as_image(path: str, lines: list) -> None:
    """PDF를 이미지로 변환하여 분석 파이프라인 (SSA → OCR → MiniCPM) 실행.

    텍스트 추출 불가능한 스캔 문서 PDF를 처리.
    """
    try:
        import fitz
        doc = fitz.open(path)
        if doc.page_count == 0:
            lines.append("⚠️ PDF has no pages.")
            doc.close()
            return

        lines.append("")
        lines.append("### 🔬 PDF → Image Pipeline (Scanned Document)")
        lines.append(f"Text extraction failed — {doc.page_count} page(s) being analyzed as image.")
        lines.append("")

        # 첫 페이지만 이미지 변환
        page = doc[0]
        pix = page.get_pixmap(dpi=200)
        img_path = path + ".png"
        pix.save(img_path)
        doc.close()

        # ── (SSA 생략: PDF 문서는 공간 분석 대상이 아님) ──
        lines.append("")

        # ── OCR 텍스트 추출 (한국어 우선) ──
        try:
            from bridge.ocr_engine import OcrEngine
            ocr = OcrEngine()
            if ocr.is_available():
                result = ocr.ocr(img_path, lang="kor", detail="full")
                md = OcrEngine.ocr_to_markdown(result)
                lines.append(md)
                lines.append("")
        except Exception as e:
            lines.append(f"OCR skipped: {e}")
            lines.append("")

        # ── MiniCPM-V 비전 분석 ──
        try:
            from bridge.vision.minicpm import describe_image, is_available
            if is_available():
                desc = describe_image(img_path,
                    "이 PDF 문서의 내용을 한국어로 자세히 읽어주세요. 모든 텍스트를 추출해주세요.")
                if desc:
                    lines.append("### 🤖 Vision Analysis (MiniCPM-V)")
                    lines.append(desc)
        except Exception:
            pass

        # 임시 이미지 정리
        try:
            os.remove(img_path)
        except Exception:
            pass

    except ImportError:
        lines.append("⚠️ PyMuPDF not installed (required for scanned PDF analysis)")
    except Exception as e:
        lines.append(f"⚠️ PDF image analysis failed: {e}")


def _check_uploaded_files_impl() -> str:
    """드랍존에 업로드된 최근 파일 목록을 확인합니다."""
    registry_path = os.path.expanduser("~/.vibezoo-uploads/latest.json")

    if not os.path.exists(registry_path):
        return "📂 아직 업로드된 파일이 없습니다."

    # 세션 시작 시간 읽기
    session_start = 0.0
    try:
        if os.path.exists(DZ_SESSION_FILE):
            with open(DZ_SESSION_FILE, 'r') as f:
                session = json.load(f)
            session_start = session.get("started_at", 0.0)
    except Exception:
        pass

    # 세션 파일이 없으면 최근 5분 이내 파일만 표시 (fallback)
    if session_start == 0.0:
        session_start = time.time() - 300

    try:
        with open(registry_path, 'r') as f:
            entries = json.load(f)

        # 세션 시작 이후 항목만 필터링 (latest.json은 ms, session은 s)
        entries = [
            e for e in entries
            if (e.get("timestamp", 0) / 1000.0) >= session_start
        ]

        if not entries:
            return "📂 현재 세션에 업로드된 파일이 없습니다. 드롭존에 파일을 업로드해주세요."

        lines = ["## 📎 최근 업로드된 파일", ""]
        for i, entry in enumerate(entries):
            path = entry.get("path", "?")
            name = entry.get("fileName", "?")
            size = entry.get("size", 0)
            mime = entry.get("mimeType", "?")

            size_str = f"{size/1024:.1f}KB" if size > 1024 else f"{size}B"
            lines.append(f"### {i+1}. {name}")
            lines.append(f"- **경로**: `{path}`")
            lines.append(f"- **크기**: {size_str}")
            lines.append(f"- **타입**: {mime}")
            lines.append("")

        if entries:
            lines.append(f"**분석 예시**: `analyze_uploaded_file(file_path='{entries[0]['path']}')`")
        return "\n".join(lines)
    except Exception as e:
        return f"⚠️ 업로드 레지스트리 읽기 실패: {e}"


def register(mcp):
    """파일 분석 도구 등록"""
    
    @mcp.tool
    def analyze_uploaded_file(file_path: str = "") -> str:
        """드롭존에 업로드된 파일을 분석하거나, 인자가 없으면 업로드된 파일 목록을 확인합니다.

        file_path를 제공하지 않으면 (기본값) 최근 업로드된 파일 목록을 반환합니다.

        파일 타입 자동 감지 → 분석 파이프라인 실행:
        - 이미지: SSA 공간 분석 → OCR 텍스트 추출 → MiniCPM-V 비전 분석
        - 코드: 내용 읽기 → 구문 분석 제안
        - 문서: PDF/DOCX 텍스트 추출

        분석 완료 후 사용자에게 "무엇을 해드릴까요?" 후속 질문을 제안합니다.

        Args:
            file_path: 업로드된 파일의 전체 경로 (생략 시 최근 업로드된 파일 목록 확인)
        Returns:
            마크다운 형식의 파일 분석 보고서 또는 업로드된 파일 목록
        """
        if not file_path:
            return _check_uploaded_files_impl()
        return analyze_file(file_path)

