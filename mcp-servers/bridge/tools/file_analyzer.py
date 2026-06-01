"""Universal File Analyzer - analyze any uploaded file by type"""
import os, json, base64
from pathlib import Path
from datetime import datetime
from bridge.utils import _markdown_header, _markdown_footer

def _get_file_info(path: str) -> dict:
    """Get file metadata"""
    p = Path(path)
    stat = p.stat()
    ext = p.suffix.lower()
    size_kb = stat.st_size / 1024
    
    # Determine type category
    image_exts = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg', '.ico'}
    doc_exts = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt'}
    text_exts = {'.txt', '.md', '.rst', '.json', '.yaml', '.yml', '.csv', '.xml', '.toml', '.ini', '.cfg'}
    code_exts = {'.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.cpp', '.c', '.h', '.hpp', 
                 '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.sh', '.bat', '.ps1'}
    archive_exts = {'.zip', '.tar', '.gz', '.rar', '.7z'}
    
    if ext in image_exts:
        ftype = "image"
    elif ext in doc_exts:
        ftype = "document"
    elif ext in text_exts:
        ftype = "text"
    elif ext in code_exts:
        ftype = "code"
    elif ext in archive_exts:
        ftype = "archive"
    else:
        ftype = "binary"
    
    return {
        "path": str(p.absolute()),
        "name": p.name,
        "ext": ext,
        "type": ftype,
        "size_kb": round(size_kb, 1),
        "size_mb": round(size_kb / 1024, 2),
        "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }

def _read_text_file(path: str, max_chars: int = 50000) -> str:
    """Read text file with encoding detection"""
    encodings = ["utf-8", "euc-kr", "cp949", "latin-1", "shift-jis"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                text = f.read(max_chars)
            if len(text) >= max_chars:
                text += f"\n\n...[truncated at {max_chars} chars]"
            return text, enc
        except UnicodeDecodeError:
            continue
    return "", "unknown"

def analyze_file(file_path: str) -> str:
    """파일 분석 메인 함수
    
    Args:
        file_path: 분석할 파일 경로
    Returns:
        마크다운 형식 분석 보고서
    """
    path = os.path.expanduser(file_path)
    if not os.path.exists(path):
        return (_markdown_header("File Error", "❌")
                + f"**File not found:** `{path}`\n"
                + _markdown_footer())
    
    info = _get_file_info(path)
    
    lines = []
    lines.append(f"## 📄 File Analysis: {info['name']}")
    lines.append("")
    lines.append(f"- **Type:** {info['type']} ({info['ext']})")
    lines.append(f"- **Size:** {info['size_kb']} KB ({info['size_mb']} MB)")
    lines.append(f"- **Path:** `{info['path']}`")
    lines.append(f"- **Modified:** {info['modified']}")
    lines.append("")
    
    # For images: include data URI for vision
    if info['type'] == 'image':
        try:
            with open(path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode()
            ext = info['ext'].lstrip('.') or 'png'
            if ext == 'jpg': ext = 'jpeg'
            data_uri = f"data:image/{ext};base64,{img_b64[:200000]}"
            lines.append(f"![uploaded image]({data_uri})")
            lines.append("")
        except Exception:
            pass
        
        # Try MiniCPM-V vision
        try:
            from bridge.vision.minicpm import describe_image, is_available
            if is_available():
                desc = describe_image(path)
                if desc:
                    lines.append("### 🤖 Vision Analysis")
                    lines.append(desc)
                    lines.append("")
        except Exception:
            pass
    
    # For text/code files: read content
    if info['type'] in ('text', 'code'):
        content, encoding = _read_text_file(path)
        lines.append(f"### 📝 Content ({encoding})")
        lines.append(f"```{info['ext'].lstrip('.')}")
        lines.append(content[:10000])
        if len(content) > 10000:
            lines.append("...[truncated]")
        lines.append("```")
        lines.append("")
    
    # For documents: try text extraction
    if info['type'] == 'document':
        if info['ext'] == '.pdf':
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(path)
                lines.append(f"### 📑 PDF Content ({doc.page_count} pages)")
                for i, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        lines.append(f"\n**Page {i+1}:**")
                        lines.append(f"```\n{text[:2000]}\n```")
                    if i >= 5:  # Max 5 pages
                        lines.append("*...more pages available*")
                        break
                doc.close()
            except ImportError:
                lines.append("⚠️ PyMuPDF not installed. Run: `pip install PyMuPDF`")
            except Exception as e:
                lines.append(f"⚠️ PDF read error: {e}")
        
        elif info['ext'] == '.docx':
            try:
                import docx
                d = docx.Document(path)
                text = "\n".join(p.text for p in d.paragraphs)
                lines.append(f"### 📝 DOCX Content ({len(d.paragraphs)} paragraphs)")
                lines.append(f"```\n{text[:5000]}\n```")
            except ImportError:
                lines.append("⚠️ python-docx not installed. Run: `pip install python-docx`")
            except Exception as e:
                lines.append(f"⚠️ DOCX read error: {e}")
    
    # For binary files: show hex preview
    if info['type'] == 'binary':
        try:
            with open(path, "rb") as f:
                data = f.read(512)
            hex_str = data.hex()
            # Format as hex dump
            dump_lines = []
            for i in range(0, len(hex_str), 32):
                addr = i // 2
                hex_part = " ".join(hex_str[j:j+2] for j in range(i, min(i+32, len(hex_str)), 2))
                ascii_part = "".join(chr(int(hex_str[j:j+2], 16)) if 32 <= int(hex_str[j:j+2], 16) < 127 else "." for j in range(i, min(i+32, len(hex_str)), 2))
                dump_lines.append(f"{addr:08x}  {hex_part:<48}  {ascii_part}")
            lines.append("### 🔤 Hex Preview (first 512 bytes)")
            lines.append("```")
            lines.extend(dump_lines[:16])
            lines.append("```")
        except Exception:
            pass
    
    # Add summary
    lines.insert(1, f"*{info['size_kb']}KB {info['type']} file analyzed*")
    
    return "\n".join(lines)

def register(mcp):
    """파일 분석 도구 등록"""
    
    @mcp.tool
    def analyze_uploaded_file(file_path: str) -> str:
        """드롭존에 업로드된 파일을 분석합니다.
        파일 타입(이미지/PDF/DOCX/TXT/코드 등)을 자동 감지하여 적절한 분석 수행.
        
        Args:
            file_path: 업로드된 파일의 전체 경로
        Returns:
            마크다운 형식의 파일 분석 보고서
        """
        return analyze_file(file_path)
